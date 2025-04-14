from docx import Document
import os
from config import TEMP_FOLDER, UPLOAD_FOLDER

def fill_contract(data, modelo_nome):
    caminho_modelo = os.path.join(UPLOAD_FOLDER, modelo_nome)
    doc = Document(caminho_modelo)

    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))

    output_path = os.path.join(TEMP_FOLDER, f"contrato_{data['nome_aluno'].replace(' ', '_')}.docx")
    doc.save(output_path)
    return output_path
