from flask import Flask, render_template, request, redirect, send_file
from docx import Document
from datetime import datetime
import sqlite3
import os
import pandas as pd
from unidecode import unidecode
import json
import shutil

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = os.path.join(app.root_path, 'modelos')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.join(app.root_path, 'data'), exist_ok=True)
os.makedirs(os.path.join(app.root_path, 'static'), exist_ok=True)

DB_PATH = os.path.join(app.root_path, 'data', 'contratos.db')

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute('''
        CREATE TABLE IF NOT EXISTS cursos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL UNIQUE
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS contratos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome_aluno TEXT,
            curso TEXT,
            forma_pagamento TEXT,
            data_criacao TEXT,
            modelo_utilizado TEXT
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS modelos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            descricao TEXT,
            categoria TEXT
        )
    ''')
    conn.commit()
    conn.close()

def get_modelos():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT id, nome, descricao, categoria FROM modelos ORDER BY descricao")
    modelos = cur.fetchall()
    conn.close()
    return modelos

def get_cursos():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT nome FROM cursos ORDER BY nome ASC")
    cursos = [row[0] for row in cur.fetchall()]
    conn.close()
    return cursos

@app.route('/')
def form():
    cursos = get_cursos()
    modelos = get_modelos()
    if os.path.exists('data/alunos.json'):
        shutil.copyfile('data/alunos.json', 'static/alunos.json')
    return render_template('form.html', cursos=cursos, modelos=modelos)

@app.route('/submit_contract', methods=['POST'])
def submit_contract():
    hoje = datetime.today()
    meses = {
        '01': 'janeiro', '02': 'fevereiro', '03': 'março', '04': 'abril',
        '05': 'maio', '06': 'junho', '07': 'julho', '08': 'agosto',
        '09': 'setembro', '10': 'outubro', '11': 'novembro', '12': 'dezembro'
    }
    dia = hoje.strftime('%d')
    mes = meses[hoje.strftime('%m')]
    ano = hoje.strftime('%Y')
    data_contrato_formatada = f'Belo Horizonte, {dia} de {mes} de {ano}'

    modelo = request.form.get('modelo', '')
    data = {key: request.form.get(key, '') for key in request.form}
    data['data_contrato'] = data_contrato_formatada
    data['contratante'] = data['nome_aluno']
    data['modelo'] = modelo

    docx_path = fill_contract(data, modelo)
    return send_file(docx_path, as_attachment=True)

def fill_contract(data, modelo_nome):
    caminho_modelo = os.path.join(app.config['UPLOAD_FOLDER'], modelo_nome)
    doc = Document(caminho_modelo)

    for para in doc.paragraphs:
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))

    output_path = os.path.join(app.root_path, 'temp_filled_contract.docx')
    doc.save(output_path)
    return output_path

@app.route('/alunos_csv', methods=['GET', 'POST'])
def upload_alunos():
    mensagem = ''
    if request.method == 'POST':
        arquivo = request.files.get('arquivo')
        if arquivo and arquivo.filename.endswith('.csv'):
            try:
                df = pd.read_csv(arquivo, encoding='utf-8', sep=',')
                df.columns = [col.strip().lower() for col in df.columns]
                alunos = []
                for _, row in df.iterrows():
                    aluno = {
                        'nome': unidecode(str(row.get('nome', ''))).strip(),
                        'cpf': str(row.get('cpf', '')).strip(),
                        'email': str(row.get('email', '')).strip(),
                        'tel_aluno': str(
                            (row.get('telefone') or row.get('telefones') or row.get('celular') or '')
                        ).split('/')[0].strip(),
                        'endereco': str(row.get('endereco', '')).strip(),
                        'curso': str(row.get('produto', '')).strip()
                    }
                    alunos.append(aluno)

                with open('data/alunos.json', 'w', encoding='utf-8') as f:
                    json.dump(alunos, f, ensure_ascii=False, indent=2)

                mensagem = 'Alunos importados com sucesso!'
            except Exception as e:
                mensagem = f'Erro no processamento: {str(e)}'
        else:
            mensagem = 'Apenas arquivos CSV são aceitos.'
    return render_template('alunos_csv.html', mensagem=mensagem)

@app.route('/modelos', methods=['GET', 'POST'])
def gerenciar_modelos():
    mensagem = ''
    if request.method == 'POST':
        arquivo = request.files.get('arquivo')
        descricao = request.form.get('descricao')
        categoria = request.form.get('categoria')

        if arquivo and arquivo.filename.endswith('.docx'):
            caminho = os.path.join(app.config['UPLOAD_FOLDER'], arquivo.filename)
            arquivo.save(caminho)

            conn = sqlite3.connect(DB_PATH)
            cur = conn.cursor()
            cur.execute('INSERT INTO modelos (nome, descricao, categoria) VALUES (?, ?, ?)',
                        (arquivo.filename, descricao, categoria))
            conn.commit()
            conn.close()
            mensagem = 'Modelo enviado com sucesso!'
        else:
            mensagem = 'Envie um arquivo .docx válido.'

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute('SELECT id, nome, descricao, categoria FROM modelos ORDER BY descricao')
    modelos = cur.fetchall()
    conn.close()

    return render_template('modelos.html', modelos=modelos, mensagem=mensagem)

@app.route('/excluir_modelo/<int:modelo_id>')
def excluir_modelo(modelo_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute('SELECT nome FROM modelos WHERE id = ?', (modelo_id,))
    resultado = cur.fetchone()

    if resultado:
        nome_arquivo = resultado[0]
        caminho_arquivo = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo)
        if os.path.exists(caminho_arquivo):
            os.remove(caminho_arquivo)
        cur.execute('DELETE FROM modelos WHERE id = ?', (modelo_id,))
        conn.commit()

    conn.close()
    return redirect('/modelos')

@app.route('/cursos', methods=['GET', 'POST'])
def gerenciar_cursos():
    mensagem = ''
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    if request.method == 'POST':
        novo = request.form.get('novo_curso', '').strip()
        if novo:
            try:
                cur.execute('INSERT INTO cursos (nome) VALUES (?)', (novo,))
                conn.commit()
                mensagem = 'Curso cadastrado com sucesso!'
            except sqlite3.IntegrityError:
                mensagem = 'Este curso já está cadastrado.'

    cur.execute('SELECT id, nome FROM cursos ORDER BY nome')
    cursos = cur.fetchall()
    conn.close()
    return render_template('cursos.html', cursos=cursos, mensagem=mensagem)

@app.route('/excluir_curso/<int:curso_id>')
def excluir_curso(curso_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute('DELETE FROM cursos WHERE id = ?', (curso_id,))
    conn.commit()
    conn.close()
    return redirect('/cursos')

if __name__ == '__main__':
    init_db()
    app.run(debug=True)
